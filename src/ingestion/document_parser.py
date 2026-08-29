"""
EduScribe AI - Document Parser Module
Extracts structured text and question patterns from PDF past papers, Word DOCX revision sets, and TeX files.
"""

import io
from pathlib import Path
from typing import Dict, Any, List, Optional
try:
    import pypdf
except ImportError:
    try:
        import PyPDF2 as pypdf
    except ImportError:
        pypdf = None

try:
    import docx
except ImportError:
    docx = None

class DocumentParser:
    @staticmethod
    def parse_pdf(file_bytes: bytes, filename: str = "document.pdf") -> Dict[str, Any]:
        """Parses PDF bytes into extracted page-by-page text."""
        global pypdf
        if pypdf is None:
            try:
                import pypdf
            except ImportError:
                try:
                    import PyPDF2 as pypdf
                except ImportError:
                    pypdf = None

        if pypdf is None:
            # Fallback: extract printable ASCII/text strings if PDF library unavailable
            text_preview = file_bytes.decode("utf-8", errors="ignore")
            clean_lines = [line.strip() for line in text_preview.splitlines() if len(line.strip()) > 3]
            fallback_text = "\n".join(clean_lines[:100]) if clean_lines else f"[PDF parsing requires 'pypdf': pip install pypdf]"
            return {
                "filename": filename,
                "file_type": "pdf",
                "page_count": 1,
                "pages": [{"page_num": 1, "text": fallback_text}],
                "full_text": fallback_text
            }

        try:
            reader = pypdf.PdfReader(io.BytesIO(file_bytes))
            pages_text = []
            for idx, page in enumerate(reader.pages):
                text = page.extract_text() or ""
                pages_text.append({"page_num": idx + 1, "text": text})

            full_text = "\n\n".join([f"--- Page {p['page_num']} ---\n{p['text']}" for p in pages_text])
            return {
                "filename": filename,
                "file_type": "pdf",
                "page_count": len(pages_text),
                "pages": pages_text,
                "full_text": full_text
            }
        except Exception as e:
            return {
                "filename": filename,
                "file_type": "pdf",
                "page_count": 0,
                "pages": [],
                "full_text": f"[Error reading PDF '{filename}': {str(e)}]"
            }

    @staticmethod
    def parse_docx(file_bytes: bytes, filename: str = "document.docx") -> Dict[str, Any]:
        """Parses Word DOCX bytes into structured paragraph text and tables."""
        try:
            doc = docx.Document(io.BytesIO(file_bytes))
            paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
            
            tables_text = []
            for table in doc.tables:
                table_rows = []
                for row in table.rows:
                    row_data = [cell.text.strip() for cell in row.cells]
                    table_rows.append(" | ".join(row_data))
                tables_text.append("\n".join(table_rows))

            full_text = "\n\n".join(paragraphs)
            if tables_text:
                full_text += "\n\n--- Tables Extracted ---\n" + "\n\n".join(tables_text)

            return {
                "filename": filename,
                "file_type": "docx",
                "paragraph_count": len(paragraphs),
                "table_count": len(tables_text),
                "full_text": full_text
            }
        except Exception:
            text = file_bytes.decode("utf-8", errors="ignore")
            return {
                "filename": filename,
                "file_type": "docx",
                "paragraph_count": len(text.splitlines()),
                "table_count": 0,
                "full_text": text
            }

    @classmethod
    def parse_file(cls, file_bytes: bytes, filename: str) -> Dict[str, Any]:
        """Routes file by extension."""
        ext = Path(filename).suffix.lower()
        if ext == ".pdf":
            return cls.parse_pdf(file_bytes, filename)
        elif ext in [".docx", ".doc"]:
            return cls.parse_docx(file_bytes, filename)
        else:
            # Assume text/tex/md/txt
            text = file_bytes.decode("utf-8", errors="ignore")
            return {
                "filename": filename,
                "file_type": ext.lstrip("."),
                "full_text": text
            }
